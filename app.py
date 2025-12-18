import os
import uuid
import json
import zipfile
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
import threading

from flask import Flask, request, jsonify, render_template, send_file, url_for
from werkzeug.utils import secure_filename
import cv2
import numpy as np
from PIL import Image
import subprocess

# OCR相关导入
try:
    from paddleocr import PaddleOCR
    import paddle
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("警告: PaddleOCR未安装，将使用模拟OCR功能")

# 文档处理导入
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，将无法生成docx文件")

# PDF处理导入
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("警告: pdf2image未安装，将无法处理PDF文件")

app = Flask(__name__)

# 配置
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'output'
app.config['TMP_FOLDER'] = 'tmp'

# 创建必要的目录
for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER'], app.config['TMP_FOLDER']]:
    os.makedirs(folder, exist_ok=True)

# 任务状态存储
tasks = {}
tasks_lock = threading.Lock()

# OCR实例（延迟加载）
ocr_instance = None
ocr_lock = threading.Lock()

def get_ocr_instance():
    """获取OCR实例，线程安全"""
    global ocr_instance
    if ocr_instance is None and OCR_AVAILABLE:
        with ocr_lock:
            if ocr_instance is None:
                try:
                    # 使用中文模型，启用方向分类器和版面分析
                    ocr_instance = PaddleOCR(
                        use_angle_cls=True,
                        lang='ch',
                        show_log=False,
                        use_gpu=False,  # CPU部署
                        max_batch_size=4
                    )
                except Exception as e:
                    print(f"OCR初始化失败: {e}")
                    return None
    return ocr_instance

def allowed_file(filename):
    """检查文件类型是否允许"""
    allowed_extensions = {'pdf', 'png', 'jpg', 'jpeg', 'gif', 'bmp', 'zip'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def extract_zip(zip_path, extract_to):
    """解压zip文件"""
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)
    
    # 获取所有图片文件
    image_files = []
    for root, dirs, files in os.walk(extract_to):
        for file in files:
            if file.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                image_files.append(os.path.join(root, file))
    
    # 按文件名排序
    image_files.sort()
    return image_files

def pdf_to_images(pdf_path, output_dir):
    """将PDF转换为图片"""
    if not PDF2IMAGE_AVAILABLE:
        raise Exception("pdf2image未安装，无法处理PDF文件")
    
    try:
        images = convert_from_path(pdf_path, dpi=200)
        image_paths = []
        for i, image in enumerate(images):
            image_path = os.path.join(output_dir, f'page_{i+1:03d}.jpg')
            image.save(image_path, 'JPEG', quality=85)
            image_paths.append(image_path)
        return image_paths
    except Exception as e:
        raise Exception(f"PDF转换失败: {e}")

def sort_vertical_text_boxes(boxes, texts, scores):
    """
    竖排文字排序算法
    按照"右→左，上→下"的古籍阅读顺序排序
    """
    if not boxes:
        return [], [], []
    
    # 计算每个框的中心点
    centers = []
    for i, box in enumerate(boxes):
        # box格式: [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
        x_coords = [point[0] for point in box]
        y_coords = [point[1] for point in box]
        center_x = sum(x_coords) / len(x_coords)
        center_y = sum(y_coords) / len(y_coords)
        centers.append((center_x, center_y, i))
    
    # 按x坐标分组（从右到左）
    # 首先按x坐标降序排列（右边优先）
    centers.sort(key=lambda x: x[0], reverse=True)
    
    # 分组并排序
    sorted_indices = []
    current_group = []
    prev_x = centers[0][0] if centers else 0
    
    for center_x, center_y, idx in centers:
        # 如果x坐标差异较大，开始新列
        if abs(center_x - prev_x) > 50:  # 50像素阈值
            # 对当前组按y坐标排序（从上到下）
            current_group.sort(key=lambda x: x[1])
            sorted_indices.extend([idx for _, _, idx in current_group])
            current_group = []
        
        current_group.append((center_x, center_y, idx))
        prev_x = center_x
    
    # 处理最后一组
    if current_group:
        current_group.sort(key=lambda x: x[1])
        sorted_indices.extend([idx for _, _, idx in current_group])
    
    # 重新排列
    sorted_boxes = [boxes[i] for i in sorted_indices]
    sorted_texts = [texts[i] for i in sorted_indices]
    sorted_scores = [scores[i] for i in sorted_indices]
    
    return sorted_boxes, sorted_texts, sorted_scores

def process_image_with_ocr(image_path, task_id, page_num):
    """使用OCR处理单张图片"""
    ocr = get_ocr_instance()
    if ocr is None:
        # 模拟OCR结果用于测试
        return ["模拟OCR文本"], [0.95]
    
    try:
        # 读取图片
        img = cv2.imread(image_path)
        if img is None:
            raise Exception(f"无法读取图片: {image_path}")
        
        # 执行OCR
        result = ocr.ocr(img, cls=True)
        
        if not result or not result[0]:
            return [], []
        
        # 提取结果
        boxes = []
        texts = []
        scores = []
        
        for line in result[0]:
            if line and len(line) >= 2:
                box = line[0]
                text_info = line[1]
                if text_info and len(text_info) >= 2:
                    texts.append(text_info[0])
                    scores.append(text_info[1])
                    boxes.append(box)
        
        # 竖排文字排序
        sorted_boxes, sorted_texts, sorted_scores = sort_vertical_text_boxes(boxes, texts, scores)
        
        return sorted_texts, sorted_scores
        
    except Exception as e:
        print(f"OCR处理失败 {image_path}: {e}")
        return [], []

def process_task(task_id, file_path, file_type):
    """处理OCR任务"""
    with tasks_lock:
        tasks[task_id]['status'] = 'processing'
        tasks[task_id]['progress'] = 0
    
    try:
        # 创建临时目录
        temp_dir = os.path.join(app.config['TMP_FOLDER'], task_id)
        os.makedirs(temp_dir, exist_ok=True)
        
        image_paths = []
        
        # 根据文件类型处理
        if file_type == 'pdf':
            # PDF转图片
            with tasks_lock:
                tasks[task_id]['message'] = 'PDF拆页中...'
            image_paths = pdf_to_images(file_path, temp_dir)
        elif file_type == 'zip':
            # 解压图片
            with tasks_lock:
                tasks[task_id]['message'] = '解压文件中...'
            image_paths = extract_zip(file_path, temp_dir)
        else:
            # 单张图片
            image_paths = [file_path]
        
        total_pages = len(image_paths)
        all_texts = []
        
        # OCR处理每页
        for i, image_path in enumerate(image_paths):
            with tasks_lock:
                tasks[task_id]['message'] = f'OCR识别第 {i+1}/{total_pages} 页...'
                tasks[task_id]['progress'] = int((i / total_pages) * 80)
            
            texts, scores = process_image_with_ocr(image_path, task_id, i+1)
            
            # 合并文本
            page_text = '\n'.join(texts) if texts else ''
            all_texts.append(page_text)
            
            # 保存单页文本
            page_output_dir = os.path.join(app.config['OUTPUT_FOLDER'], task_id)
            os.makedirs(page_output_dir, exist_ok=True)
            
            page_txt_path = os.path.join(page_output_dir, f'single_page_{i+1:03d}_hengpai.txt')
            with open(page_txt_path, 'w', encoding='utf-8') as f:
                f.write(page_text)
        
        # 生成完整文本
        with tasks_lock:
            tasks[task_id]['message'] = '生成文档中...'
            tasks[task_id]['progress'] = 90
        
        full_text = '\n\n'.join(all_texts)
        
        # 保存完整文本
        full_txt_path = os.path.join(app.config['OUTPUT_FOLDER'], task_id, 'full_book.txt')
        with open(full_txt_path, 'w', encoding='utf-8') as f:
            f.write(full_text)
        
        # 生成Word文档
        if DOCX_AVAILABLE:
            docx_path = os.path.join(app.config['OUTPUT_FOLDER'], task_id, 'full_book.docx')
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('古籍OCR识别结果', 0)
            
            # 添加段落
            for page_text in all_texts:
                if page_text.strip():
                    paragraphs = page_text.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            p = doc.add_paragraph(para_text)
                            p.paragraph_format.space_after = Pt(6)
                    # 添加分页符
                    doc.add_page_break()
            
            doc.save(docx_path)
        
        # 打包成zip
        with tasks_lock:
            tasks[task_id]['message'] = '打包文件中...'
            tasks[task_id]['progress'] = 95
        
        zip_path = os.path.join(app.config['OUTPUT_FOLDER'], f'{task_id}.zip')
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            output_dir = os.path.join(app.config['OUTPUT_FOLDER'], task_id)
            for root, dirs, files in os.walk(output_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, output_dir)
                    zipf.write(file_path, arcname)
        
        # 更新任务状态
        with tasks_lock:
            tasks[task_id]['status'] = 'completed'
            tasks[task_id]['progress'] = 100
            tasks[task_id]['message'] = '处理完成'
            tasks[task_id]['result_url'] = url_for('download', task_id=task_id, _external=True)
        
        # 清理临时文件
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
        
    except Exception as e:
        with tasks_lock:
            tasks[task_id]['status'] = 'failed'
            tasks[task_id]['message'] = f'处理失败: {str(e)}'
            tasks[task_id]['error'] = str(e)
        print(f"任务 {task_id} 失败: {e}")

@app.route('/')
def index():
    """主页"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    """文件上传接口"""
    if 'file' not in request.files:
        return jsonify({'error': '未找到文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': '不支持的文件格式'}), 400
    
    # 生成任务ID
    task_id = str(uuid.uuid4())
    
    # 保存文件
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task_id}_{filename}")
    file.save(file_path)
    
    # 确定文件类型
    file_ext = filename.rsplit('.', 1)[1].lower()
    if file_ext == 'pdf':
        file_type = 'pdf'
    elif file_ext == 'zip':
        file_type = 'zip'
    else:
        file_type = 'image'
    
    # 初始化任务状态
    with tasks_lock:
        tasks[task_id] = {
            'status': 'pending',
            'progress': 0,
            'message': '等待处理',
            'filename': filename,
            'file_type': file_type,
            'created_at': datetime.now().isoformat()
        }
    
    # 异步处理任务
    executor = ThreadPoolExecutor(max_workers=2)
    executor.submit(process_task, task_id, file_path, file_type)
    
    return jsonify({
        'task_id': task_id,
        'status_url': url_for('status', task_id=task_id, _external=True)
    })

@app.route('/status/<task_id>')
def status(task_id):
    """查询任务状态"""
    with tasks_lock:
        if task_id not in tasks:
            return jsonify({'error': '任务不存在'}), 404
        
        task_info = tasks[task_id].copy()
        return jsonify(task_info)

@app.route('/download/<task_id>')
def download(task_id):
    """下载结果文件"""
    zip_path = os.path.join(app.config['OUTPUT_FOLDER'], f'{task_id}.zip')
    
    if not os.path.exists(zip_path):
        return jsonify({'error': '文件不存在'}), 404
    
    with tasks_lock:
        filename = tasks.get(task_id, {}).get('filename', 'result')
    
    download_filename = f"{filename.rsplit('.', 1)[0]}_OCR结果.zip"
    
    return send_file(zip_path, as_attachment=True, download_name=download_filename)

@app.route('/preview/<task_id>')
def preview(task_id):
    """预览识别结果"""
    txt_path = os.path.join(app.config['OUTPUT_FOLDER'], task_id, 'full_book.txt')
    
    if not os.path.exists(txt_path):
        return jsonify({'error': '结果不存在'}), 404
    
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 返回前1000字符作为预览
    preview_text = content[:1000]
    if len(content) > 1000:
        preview_text += '...'
    
    return jsonify({
        'preview': preview_text,
        'total_length': len(content)
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)